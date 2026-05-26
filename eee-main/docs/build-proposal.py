"""
Convert npic-qc-proposal-v1.md → npic-qc-proposal-v1.docx
using python-docx with proper Chinese-friendly formatting.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re, os

# ── helpers ────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    """Set table cell background colour (hex without #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if val:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'), 'single')
            el.set(qn('w:sz'), '4')
            el.set(qn('w:color'), val)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para_space(para, before=0, after=0):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)

def add_rule(doc):
    """Add a horizontal rule paragraph."""
    p = doc.add_paragraph()
    para_space(p, before=4, after=4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), 'D1D5DB')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_run_inline(para, text, bold=False, italic=False, code=False, color=None):
    run = para.add_run(text)
    run.bold   = bold
    run.italic = italic
    if code:
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x1f, 0x29, 0x37)
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

# ── inline markdown parser ─────────────────────────────────────────────────

INLINE_RE = re.compile(
    r'(`[^`]+`)'           # inline code
    r'|(\*\*[^*]+\*\*)'    # bold
    r'|(\*[^*]+\*)'        # italic
)

def render_inline(para, text):
    pos = 0
    for m in INLINE_RE.finditer(text):
        if m.start() > pos:
            para.add_run(text[pos:m.start()])
        chunk = m.group(0)
        if chunk.startswith('`'):
            add_run_inline(para, chunk[1:-1], code=True)
        elif chunk.startswith('**'):
            add_run_inline(para, chunk[2:-2], bold=True)
        else:
            add_run_inline(para, chunk[1:-1], italic=True)
        pos = m.end()
    if pos < len(text):
        para.add_run(text[pos:])

# ── document builder ───────────────────────────────────────────────────────

def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # Default paragraph font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    md_path = os.path.join(os.path.dirname(__file__), 'npic-qc-proposal-v1.md')
    lines = open(md_path, encoding='utf-8').readlines()

    i = 0
    while i < len(lines):
        raw = lines[i].rstrip('\n')
        stripped = raw.strip()

        # ── blank / HR ──────────────────────────────────────────────────
        if not stripped:
            i += 1
            continue

        if stripped.startswith('---'):
            add_rule(doc)
            i += 1
            continue

        # ── headings ────────────────────────────────────────────────────
        m = re.match(r'^(#{1,4})\s+(.*)', stripped)
        if m:
            level = len(m.group(1))
            text  = m.group(2)
            if level == 1:
                p = doc.add_heading('', level=0)
                r = p.add_run(text)
                r.font.size = Pt(22)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)
                para_space(p, before=0, after=10)
            elif level == 2:
                p = doc.add_heading('', level=1)
                r = p.add_run(text)
                r.font.size = Pt(14)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x1d, 0x4e, 0xd8)
                para_space(p, before=18, after=4)
            elif level == 3:
                p = doc.add_heading('', level=2)
                r = p.add_run(text)
                r.font.size = Pt(11.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x37, 0x41, 0x51)
                para_space(p, before=12, after=2)
            else:
                p = doc.add_heading('', level=3)
                r = p.add_run(text)
                r.font.size = Pt(10.5)
                r.font.bold = True
                r.font.color.rgb = RGBColor(0x4b, 0x55, 0x63)
                para_space(p, before=8, after=2)
            i += 1
            continue

        # ── code block ──────────────────────────────────────────────────
        if stripped.startswith('```'):
            i += 1
            code_lines = []
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i].rstrip('\n'))
                i += 1
            i += 1  # skip closing ```
            p = doc.add_paragraph()
            para_space(p, before=4, after=4)
            pf = p.paragraph_format
            pf.left_indent = Cm(0.5)
            # grey background via shading on the paragraph
            pPr = p._p.get_or_add_pPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), '1E293B')
            pPr.append(shd)
            r = p.add_run('\n'.join(code_lines))
            r.font.name = 'Courier New'
            r.font.size = Pt(8.5)
            r.font.color.rgb = RGBColor(0xe2, 0xe8, 0xf0)
            continue

        # ── blockquote ──────────────────────────────────────────────────
        if stripped.startswith('>'):
            text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent  = Cm(0.8)
            pf.right_indent = Cm(0.8)
            para_space(p, before=4, after=4)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            left = OxmlElement('w:left')
            left.set(qn('w:val'), 'single')
            left.set(qn('w:sz'), '12')
            left.set(qn('w:color'), 'D1D5DB')
            pBdr.append(left)
            pPr.append(pBdr)
            r = p.add_run(text)
            r.font.color.rgb = RGBColor(0x6b, 0x72, 0x80)
            r.font.italic = True
            i += 1
            continue

        # ── table ───────────────────────────────────────────────────────
        if stripped.startswith('|'):
            # collect all table lines
            tbl_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                tbl_lines.append(lines[i].strip())
                i += 1
            # parse
            def parse_row(line):
                return [c.strip() for c in line.strip('|').split('|')]

            rows = [parse_row(l) for l in tbl_lines if not re.match(r'^\|[\s\-|]+\|$', l)]
            if not rows:
                continue
            cols = len(rows[0])
            tbl = doc.add_table(rows=len(rows), cols=cols)
            tbl.style = 'Table Grid'
            tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
            for r_idx, row_data in enumerate(rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = tbl.cell(r_idx, c_idx)
                    # strip inline markdown
                    clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', cell_text)
                    clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
                    clean = re.sub(r'`([^`]+)`', r'\1', clean)
                    p = cell.paragraphs[0]
                    p.text = clean
                    p.paragraph_format.space_before = Pt(3)
                    p.paragraph_format.space_after  = Pt(3)
                    if r_idx == 0:
                        p.runs[0].bold = True if p.runs else None
                        set_cell_bg(cell, 'EFF6FF')
                    for run in p.runs:
                        run.font.size = Pt(9.5)
            doc.add_paragraph()  # spacing after table
            continue

        # ── unordered list ───────────────────────────────────────────────
        if re.match(r'^[-*]\s', stripped):
            text = re.sub(r'^[-*]\s+', '', stripped)
            p = doc.add_paragraph(style='List Bullet')
            para_space(p, before=1, after=1)
            p.paragraph_format.left_indent = Cm(0.6)
            render_inline(p, text)
            i += 1
            continue

        # ── ordered list ─────────────────────────────────────────────────
        if re.match(r'^\d+\.\s', stripped):
            text = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            para_space(p, before=1, after=1)
            p.paragraph_format.left_indent = Cm(0.6)
            render_inline(p, text)
            i += 1
            continue

        # ── cover meta lines (bold: value) ───────────────────────────────
        if re.match(r'^\*\*[^*]+\*\*', stripped) and '：' in stripped:
            p = doc.add_paragraph()
            para_space(p, before=2, after=2)
            render_inline(p, stripped)
            i += 1
            continue

        # ── normal paragraph ─────────────────────────────────────────────
        p = doc.add_paragraph()
        para_space(p, before=2, after=4)
        p.paragraph_format.line_spacing = Pt(18)
        render_inline(p, stripped)
        i += 1

    out_path = os.path.join(os.path.dirname(__file__), 'npic-qc-proposal-v1.docx')
    doc.save(out_path)
    print(f'Saved: {out_path}')

if __name__ == '__main__':
    build()
